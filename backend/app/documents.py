from __future__ import annotations

import base64
import io
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

import pymupdf
from docx import Document

SUPPORTED_RESUME_TYPES = {
    ".pdf": "pdf",
    ".docx": "docx",
}


def safe_filename(filename: str) -> str:
    name = Path(filename).name
    name = re.sub(r"[^\w\-. ]+", "_", name, flags=re.UNICODE).strip()
    return name or "uploaded-document"


def parse_pdf(data: bytes) -> dict[str, Any]:
    document = pymupdf.open(stream=data, filetype="pdf")
    pages: list[dict[str, Any]] = []
    text_parts: list[str] = []
    image_count = 0
    layout_profile: dict[str, Any] = {}
    for index, page in enumerate(document):
        text = page.get_text("text", sort=True).strip()
        images = page.get_images(full=True)
        image_count += len(images)
        text_parts.append(text)
        if index == 0:
            spans = [span for block in page.get_text("dict", sort=True).get("blocks", []) for line in block.get("lines", []) for span in line.get("spans", []) if span.get("text", "").strip()]
            size_counts = Counter()
            font_counts = Counter()
            for span in spans:
                weight = len(span["text"].strip())
                size_counts[round(float(span["size"]), 2)] += weight
                font_counts[str(span["font"])] += weight
            title_sizes = [float(span["size"]) for span in spans if span["text"].strip() in {"教育经历", "实习经历", "在校经历", "相关技能", "专业技能", "项目经历", "工作经历"}]
            boxes = [span["bbox"] for span in spans]
            dominant_size = size_counts.most_common(1)[0][0] if size_counts else 10.5
            max_size = max((float(span["size"]) for span in spans), default=20.0)
            inferred_top = round(max(6.0, min((box[1] for box in boxes), default=36.0) * 25.4 / 72 - 2.0), 1)
            layout_profile = {
                "font_size": dominant_size,
                "font_family": font_counts.most_common(1)[0][0] if font_counts else "Microsoft YaHei",
                "name_font_size": round(max_size, 2),
                "section_font_size": round(Counter(round(size, 2) for size in title_sizes).most_common(1)[0][0], 2) if title_sizes else 12.0,
                "margin_left": round(min((box[0] for box in boxes), default=39.0) * 25.4 / 72, 1),
                "margin_right": round((page.rect.width - max((box[2] for box in boxes), default=550.0)) * 25.4 / 72, 1),
                "margin_top": inferred_top,
                "margin_bottom": 6.0,
            }
        pages.append({
            "page": index + 1,
            "width": round(page.rect.width, 2),
            "height": round(page.rect.height, 2),
            "text": text,
            "image_count": len(images),
        })
    full_text = "\n\n".join(part for part in text_parts if part)
    assets: list[dict[str, Any]] = []
    for page_index, page in enumerate(document):
        for image in page.get_images(full=True):
            xref = image[0]
            extracted = document.extract_image(xref)
            if not extracted:
                continue
            ext = extracted.get("ext", "png")
            assets.append({
                "name": f"pdf-image-{xref}.{ext}",
                "mime": f"image/{ext}",
                "page": page_index + 1,
                "width": extracted.get("width"),
                "height": extracted.get("height"),
                "rects": [list(map(float, rect)) for rect in page.get_image_rects(xref)],
                "data_url": f"data:image/{ext};base64,{base64.b64encode(extracted['image']).decode('ascii')}",
            })
    if assets and assets[0].get("rects") and pages:
        x0, y0, x1, y1 = assets[0]["rects"][0]
        layout_profile["photo_rect_pct"] = {
            "left": round(x0 / pages[0]["width"] * 100, 3),
            "top": round(y0 / pages[0]["height"] * 100, 3),
            "width": round((x1 - x0) / pages[0]["width"] * 100, 3),
            "height": round((y1 - y0) / pages[0]["height"] * 100, 3),
        }
    warnings: list[str] = []
    if not full_text:
        warnings.append("PDF 没有可直接提取的文本，后续需要 OCR。")
    return {
        "format": "pdf",
        "page_count": len(document),
        "text": full_text,
        "pages": pages,
        "image_count": image_count,
        "assets": assets,
        "layout_profile": layout_profile,
        "warnings": warnings,
    }


def _docx_images(data: bytes) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/"):
                images.append({"name": Path(name).name, "size": archive.getinfo(name).file_size})
    return images


def parse_docx(data: bytes) -> dict[str, Any]:
    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    table_text = ["\t".join(cell for cell in row) for table in tables for row in table]
    text = "\n".join(paragraphs + table_text)
    section = document.sections[0] if document.sections else None
    page = None
    if section:
        page = {
            "width": round(section.page_width / 914400, 2),
            "height": round(section.page_height / 914400, 2),
            "margin_left": round(section.left_margin / 914400, 2),
            "margin_right": round(section.right_margin / 914400, 2),
            "margin_top": round(section.top_margin / 914400, 2),
            "margin_bottom": round(section.bottom_margin / 914400, 2),
        }
    warnings = []
    if any("w:drawing" in paragraph._p.xml for paragraph in document.paragraphs):
        warnings.append("文档包含浮动对象，当前阶段仅记录图片资源，复杂定位将在模板阶段处理。")
    assets = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/"):
                continue
            raw = archive.read(name)
            ext = Path(name).suffix.lstrip(".") or "png"
            assets.append({"name": Path(name).name, "mime": f"image/{ext}", "data_url": f"data:image/{ext};base64,{base64.b64encode(raw).decode('ascii')}"})
    return {
        "format": "docx",
        "page_count": None,
        "text": text,
        "paragraphs": paragraphs,
        "tables": tables,
        "images": _docx_images(data),
        "assets": assets,
        "page": page,
        "warnings": warnings,
    }


def parse_text_file(filename: str, data: bytes) -> dict[str, Any]:
    text = data.decode('utf-8-sig', errors='replace').strip()
    return {"format": "text", "text": text, "warnings": [] if text else ["文件没有可读取的文字。"]}


def parse_document(filename: str, data: bytes) -> dict[str, Any]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(data)
    if suffix == ".docx":
        return parse_docx(data)
    if suffix in {".txt", ".md"}:
        return parse_text_file(filename, data)
    if suffix == ".doc":
        raise ValueError("暂不直接解析 DOC，请先转换为 DOCX。")
    raise ValueError("目前支持 PDF 和 DOCX 文件。")


def to_json(value: dict[str, Any]) -> str:
    return json.dumps(value, ensure_ascii=False)
